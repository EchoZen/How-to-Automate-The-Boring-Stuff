import docx, os
os.chdir(r'C:\Users\zchen\OneDrive\Documents\Python stuff\Automate The Boring Stuff')
d= docx.Document('Biochemistry I Lecture 3.docx')
print(d.paragraphs[100].text)
p= d.paragraphs[100]
# run is like either a new sentence of para or
# can see and change style using runs
print(p.runs)
p.runs[0].underline= True
d.paragraphs[99].runs[0].bold= True
d.paragraphs[97].style= 'Heading 1'
d.save('Biochem I le3 revised.docx')


# Create a new word doc
d= docx.Document()
p1= d.add_paragraph("Hello, this is a paragraph.")
p2= d.add_paragraph('This is another paragraph.')
p1.add_run('2nd line of 1st paragraph.')
p1.runs[1].bold= True
d.save('MynewWordDoc.docx')


# Get the text in Word as a string
def getText(filename):
    doc= docx.Document(filename)
    fullText=[]
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

stringdoc= getText('Biochemistry I Lecture 3.docx')
print(stringdoc)
